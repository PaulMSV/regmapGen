
import argparse
import re
from docx import Document
from docx.shared import Cm
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# Функция настройки ширины колонок
def change_column_width(doc, column_widths):
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell.width = Cm(column_widths[j])

# Функция настройки высоты строк
def change_row_height(table, row_index, height_cm):
    row = table.rows[row_index]
    row.height = Cm(height_cm)

# Функция изменения шрифта для таблиц и для текста
def change_font_size(doc, text_font_size, table_font_size):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(table_font_size)
    for paragraph in doc.paragraphs:
        if "Смещение адреса" in paragraph.text or "Значение при сбросе" in paragraph.text:
            for run in paragraph.runs:
                run.font.size = Pt(text_font_size)

# Функция изменения стиля для названия таблиц
def apply_table_caption_style(doc):
    for paragraph in doc.paragraphs:
        if "Описание полей регистра" in paragraph.text:
            paragraph.style = 'Нумерация таблиц 2 ур.' 

# Функция выравниевания первой строки заголовка таблиц по центру
def align_first_row_in_table_to_center(table):
    for cell in table.rows[0].cells:
        cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

# Функция добавления двойной линии в первой строке таблицы (для нижней границы)
def add_double_borders_to_first_row(table):
    for cell in table.rows[0].cells:
        tcPr = cell._element.get_or_add_tcPr()
        tcBorders = tcPr.xpath('.//w:tcBorders')
        if not tcBorders:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        else:
            tcBorders = tcBorders[0]
        for edge in ('top', 'bottom', 'start', 'end'):
            border = OxmlElement(f'w:{edge}')
            if edge == 'bottom':
                border.set(qn('w:val'), 'double')
            tcBorders.append(border)

# Функция добавления двойной линии в второй строки таблицы (для верхней границы)
def add_double_borders_to_second_row(table):
    for cell in table.rows[1].cells:
        tcPr = cell._element.get_or_add_tcPr()
        tcBorders = tcPr.xpath('.//w:tcBorders')
        if not tcBorders:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        else:
            tcBorders = tcBorders[0]
        for edge in ('top', 'bottom', 'start', 'end'):
            border = OxmlElement(f'w:{edge}')
            if edge == 'top':
                border.set(qn('w:val'), 'double')
            tcBorders.append(border)

# Функция добавления переносов строки перед заданными фрагментами текста
def add_newline_before_pattern(doc, pattern_str):
    pattern = re.compile(pattern_str)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        # Find all matches of the pattern
                        matches = list(pattern.finditer(run.text))
                        if not matches:
                            continue

                        # Split the text by the pattern and add new lines
                        new_run_texts = []
                        start = 0
                        for match in matches:
                            # Add text before the match
                            new_run_texts.append(run.text[start:match.start()])
                            # Add the new line before the match
                            new_run_texts.append('\n' + match.group(0))
                            start = match.end()
                        # Add remaining text after the last match
                        new_run_texts.append(run.text[start:])

                        # Update the run text
                        run.text = ''.join(new_run_texts)

# Функция делает текст жирным только в первой строке каждой ячейки описания, исключая заголовок таблицы
def make_first_line_bold_in_last_column(doc):
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            if i == 0:  # Пропускаем первую строку, так как это заголовки
                continue
            last_cell = row.cells[-1]  # Последняя ячейка в строке
            for paragraph in last_cell.paragraphs:
                if paragraph.text:  # Убеждаемся, что есть текст для стилизации
                    runs = paragraph.runs
                    if runs:  # Если в параграфе есть элементы
                        # Делим текст на первую строку и оставшийся текст
                        text_parts = runs[0].text.split('\n', 1)
                        first_line_text = text_parts[0]
                        # Обновляем текст первого run до первой строки и делаем его жирным
                        runs[0].text = first_line_text
                        runs[0].bold = True
                        # Если после первой строки есть ещё текст, добавляем его как новый run
                        if len(text_parts) > 1:
                            rest_of_text = text_parts[1]
                            # Создаём новый run, если есть текст после первой строки
                            new_run = paragraph.add_run('\n' + rest_of_text)
                            new_run.bold = False


def main(input_file, output_file):
    # Открываем документ
    doc = Document(input_file)

    # Напечатать в консоль название всех стилей в документе
    #for style in doc.styles:
    #    print(style.name)

    # Указываем новую ширину столбцов для каждой колонки в сантиметрах
    column_widths = [3.5, 1.2, 1.0, 1.3, 9.5]  # Пример для пяти колонок, каждая указана в сантиметрах

    # Указываем новый размер шрифта (в пунктах) для текста
    text_font_size = 10
    # Указываем новый размер шрифта (в пунктах) для таблиц
    table_font_size = 9  

    # Создаем паттерн для поиска
    pattern = r"1.b|2.b|3.b|4.b"

    ###################### Основная часть #######################

    # Изменяем ширину столбцов
    change_column_width(doc, column_widths)

    # Применяем стиль к каждой таблице
    for table in doc.tables:
        table.style = 'Table Grid'
        change_row_height(table, 0, 1.0)
        align_first_row_in_table_to_center(table)
        add_double_borders_to_first_row(table)
        add_double_borders_to_second_row(table)

    # Применяем стиль к строке "Описание полей регистра"
    apply_table_caption_style(doc)

    # Добавляем переносы строки перед "1'b0" и т.п.
    add_newline_before_pattern(doc, pattern)

    # Делаем текст жирным только в первой строке каждой ячейки последней колонки
    make_first_line_bold_in_last_column(doc)

    # Изменяем размер шрифта
    change_font_size(doc, text_font_size, table_font_size)

    # Сохраняем изменения
    doc.save(output_file)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Modify Word document with specified parameters.')
    parser.add_argument('-i', '--input', help='Input Word document file path', required=True)
    parser.add_argument('-o', '--output', help='Output Word document file path', required=True)
    args = parser.parse_args()
    main(args.input, args.output)